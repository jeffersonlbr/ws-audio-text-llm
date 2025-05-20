import os
import logging
import shutil
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from datetime import datetime

# Configuração de logging
logger = logging.getLogger(__name__)


def create_docx(dialogue, output_path, project_name, description, meta_info):
    """
    Cria um arquivo DOCX formatado com o diálogo, usando o template fornecido
    e adicionando as informações solicitadas na página inicial
    """
    try:
        # Usar o template na raiz do projeto
        template_path = "depoimento_template.docx"
        
        # Verificação do template
        if not os.path.exists(template_path):
            logger.error(f"Template DOCX não encontrado em: {template_path}")
            raise FileNotFoundError(f"Template DOCX não encontrado em: {template_path}")
            
        logger.info(f"Template encontrado em: {template_path}")
        
        # Verificar se o arquivo não está corrompido
        if os.path.getsize(template_path) == 0:
            logger.error(f"Template DOCX está vazio ou corrompido: {template_path}")
            raise ValueError("Template DOCX está vazio ou corrompido")
            
        # Verificar permissões de leitura
        if not os.access(template_path, os.R_OK):
            logger.error(f"Sem permissão para ler o template em: {template_path}")
            raise PermissionError(f"Sem permissão para ler o template em: {template_path}")
            
        # Garantir que o texto esteja em UTF-8
        if isinstance(dialogue, list):
            dialogue = [text.encode('utf-8').decode('utf-8') if isinstance(text, str) else str(text) for text in dialogue]
        elif isinstance(dialogue, str):
            dialogue = dialogue.encode('utf-8').decode('utf-8')
        
        # Verificar se dialogue é uma lista
        if isinstance(dialogue, str):
            dialogue = dialogue.split('\n\n')
        
        # Usar o template como base
        use_template = os.path.exists(template_path)
        
        if use_template:
            try:
                # Criar uma cópia do template
                shutil.copy2(template_path, output_path)
                doc = Document(output_path)
            except Exception as e:
                logger.warning(f"Erro ao usar template: {e}. Criando documento do zero.")
                use_template = False
                doc = Document()
        else:
            # Se o template não existir, criar documento do zero
            doc = Document()
        
        # Obter a data atual para o cabeçalho
        creation_date = datetime.now().strftime("%d/%m/%Y") if not meta_info else meta_info.get('created_at', datetime.now().strftime("%d/%m/%Y"))
        filename = "" if not meta_info else meta_info.get('filename', '')
        duration = "" if not meta_info else meta_info.get('audio_duration', '')
        speakers_count = 0 if not meta_info else meta_info.get('speakers_count', 0)
        
        if use_template:
            try:
                # Limpar falas de exemplo do template
                paragraphs_to_remove = []
                for i, paragraph in enumerate(doc.paragraphs):
                    if i < len(doc.paragraphs) and ("[Speaker" in paragraph.text or "Speaker" in paragraph.text):
                        paragraphs_to_remove.append(paragraph)
                
                for paragraph in paragraphs_to_remove:
                    try:
                        p = paragraph._element
                        p.getparent().remove(p)
                    except Exception as e:
                        logger.warning(f"Erro ao remover parágrafo: {e}")
                
                # Inserir os metadados no início do documento (após o título)
                # Método mais robusto que não depende de encontrar parágrafo específico
                idx = 0
                # Encontrar o primeiro parágrafo (título)
                for i, p in enumerate(doc.paragraphs):
                    if p.text and "Depoimento Regime Disciplinar" in p.text:
                        idx = i + 1
                        break
                
                # Inserir metadados após o título
                p = doc.add_paragraph()
                run = p.add_run(f"Título: {project_name}")
                run.font.name = "Inter"
                run.font.size = Pt(10)
                
                p = doc.add_paragraph()
                run = p.add_run(f"Descrição: {description}")
                run.font.name = "Inter"
                run.font.size = Pt(10)
                
                p = doc.add_paragraph()
                run = p.add_run(f"Data: {creation_date}")
                run.font.name = "Inter"
                run.font.size = Pt(10)
                
                p = doc.add_paragraph()
                run = p.add_run(f"Duração do áudio: {duration}")
                run.font.name = "Inter"
                run.font.size = Pt(10)
                
                p = doc.add_paragraph()
                run = p.add_run(f"Quantidade de Speakers: {speakers_count}")
                run.font.name = "Inter"
                run.font.size = Pt(10)
                
                # Adicionar um título para a seção de transcrição
                heading_found = False
                for i, p in enumerate(doc.paragraphs):
                    if p.text and p.text == "Depoimento" and "Regime" not in p.text:
                        heading_found = True
                        break
                
                if not heading_found:
                    doc.add_heading("Depoimento", level=1)
            
            except Exception as e:
                logger.error(f"Erro ao configurar documento a partir do template: {e}")
                # Falhar para o método sem template
                use_template = False
                doc = Document()
                doc.add_heading("Depoimento Regime Disciplinar", 0)
        
        # Se estamos usando o método sem template ou ocorreu erro no template
        if not use_template:
            # Criar documento do zero
            doc = Document()
            doc.add_heading("Depoimento Regime Disciplinar", 0)
            
            # Adicionar metadados
            p = doc.add_paragraph("")
            run = p.add_run(f"Título: {project_name}")
            run.font.name = "Inter"
            run.font.size = Pt(10)
            
            p = doc.add_paragraph("")
            run = p.add_run(f"Descrição: {description}")
            run.font.name = "Inter"
            run.font.size = Pt(10)
            
            p = doc.add_paragraph("")
            run = p.add_run(f"Data: {creation_date}")
            run.font.name = "Inter"
            run.font.size = Pt(10)
            
            p = doc.add_paragraph("")
            run = p.add_run(f"Duração do áudio: {duration}")
            run.font.name = "Inter"
            run.font.size = Pt(10)
            
            p = doc.add_paragraph("")
            run = p.add_run(f"Quantidade de Speakers: {speakers_count}")
            run.font.name = "Inter"
            run.font.size = Pt(10)
            
            # Adicionar título de Depoimento
            doc.add_heading("Depoimento", 1)
        
        # Adicionar cada bloco de diálogo (comum a ambos os métodos)
        for block in dialogue:
            if not block.strip():
                continue
                
            parts = block.split(':', 1) if ':' in block else [None, block]
            if len(parts) == 2 and parts[0] is not None:
                speaker, text = parts
                p = doc.add_paragraph()
                speaker_run = p.add_run(f"{speaker}:")
                speaker_run.bold = True
                speaker_run.underline = True
                p.add_run(f" {text.strip()}")
            else:
                # Caso não consiga separar, adiciona o texto completo
                doc.add_paragraph(block)
        
        # Salvar o documento
        doc.save(output_path)
        logger.info(f"DOCX gerado com sucesso: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Erro ao criar DOCX: {e}")
        # Tentar um método de fallback mais simples
        try:
            doc = Document()
            doc.add_heading("Depoimento Regime Disciplinar", 0)
            doc.add_paragraph(f"Título: {project_name}")
            doc.add_paragraph(f"Descrição: {description}")
            
            # Adicionar os diálogos de forma simples
            doc.add_heading("Depoimento", 1)
            if isinstance(dialogue, list):
                for block in dialogue:
                    doc.add_paragraph(block)
            elif isinstance(dialogue, str):
                for block in dialogue.split('\n\n'):
                    doc.add_paragraph(block)
            
            doc.save(output_path)
            logger.info(f"DOCX gerado com sucesso (método fallback): {output_path}")
            return output_path
        except Exception as fallback_error:
            logger.error(f"Falha no fallback para criar DOCX: {fallback_error}")
            return None

def create_txt(text, output_file):
    """
    Cria um arquivo de texto simples
    """
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        logger.info(f"TXT gerado com sucesso: {output_file}")
        return output_file
    except Exception as e:
        logger.error(f"Erro ao criar TXT: {e}")
        return None

def create_transcript_folder(trans_id, project_name, base_dir="output"):
    """
    Cria uma pasta individual para a transcrição e retorna o caminho
    """
    try:
        # Criar nome da pasta baseado no ID e nome do projeto
        safe_project_name = project_name.replace(' ', '_').replace('/', '_')
        folder_name = f"{safe_project_name}_{trans_id[:8]}"
        
        # Caminho completo da pasta
        folder_path = os.path.join(base_dir, folder_name)
        
        # Criar pasta se não existir
        os.makedirs(folder_path, exist_ok=True)
        
        logger.info(f"Pasta de transcrição criada: {folder_path}")
        return folder_path
    except Exception as e:
        logger.error(f"Erro ao criar pasta de transcrição: {e}")
        # Fallback para um nome mais simples
        fallback_path = os.path.join(base_dir, f"transcript_{trans_id[:8]}")
        os.makedirs(fallback_path, exist_ok=True)
        return fallback_path

def format_file_size(size_bytes):
    """
    Formata o tamanho do arquivo para exibição
    """
    try:
        # Converter bytes para formato legível
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} TB"
    except Exception:
        return "Desconhecido"

def format_timestamp(timestamp):
    """
    Formata um timestamp para exibição
    """
    try:
        if not timestamp:
            return "Data desconhecida"
        
        # Se for uma string de data do SQLite
        if isinstance(timestamp, str):
            from datetime import datetime
            try:
                dt = datetime.strptime(timestamp, "%Y-%m-%d %H:%M:%S")
                return dt.strftime("%d/%m/%Y %H:%M")
            except ValueError:
                return timestamp
        
        # Se for um objeto datetime
        if hasattr(timestamp, 'strftime'):
            return timestamp.strftime("%d/%m/%Y %H:%M")
            
        return str(timestamp)
    except Exception:
        return "Data desconhecida"
    
def format_duration(seconds):
    """
    Formata duração em segundos para exibição (hh:mm:ss)
    """
    try:
        seconds = int(seconds)
        
        if seconds < 60:
            return f"{seconds} seg"
        
        minutes = seconds // 60
        remaining_seconds = seconds % 60
        
        if minutes < 60:
            return f"{minutes}m {remaining_seconds}s"
        
        hours = minutes // 60
        remaining_minutes = minutes % 60
        
        return f"{hours}h {remaining_minutes}m {remaining_seconds}s"
    except Exception:
        return "Desconhecido"

def count_words(text):
    """
    Conta o número de palavras em um texto
    """
    try:
        words = text.split()
        return len(words)
    except Exception:
        return 0